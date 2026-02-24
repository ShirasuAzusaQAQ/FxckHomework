from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


QUESTION_RE = re.compile(r"^\s*(?:第\s*)?(\d+)\s*(?:题)?\s*[\.\、:：）\)]\s*(.+)?$")


@dataclass
class QuestionItem:
    qid: str
    text: str
    paragraph: Paragraph


def _iter_paragraphs(doc: _Document) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def extract_questions(doc: _Document) -> List[QuestionItem]:
    items: List[QuestionItem] = []
    seen = set()
    for p in _iter_paragraphs(doc):
        line = p.text.strip()
        if not line:
            continue
        match = QUESTION_RE.match(line)
        if not match:
            continue
        qid = str(int(match.group(1)))
        if (id(p), qid) in seen:
            continue
        seen.add((id(p), qid))
        items.append(QuestionItem(qid=qid, text=line, paragraph=p))
    return items


def _replace_placeholders(text: str, qid: str, answer: str) -> Tuple[str, bool]:
    placeholders = [f"{{{{Q{qid}}}}}", f"{{{{q{qid}}}}}", f"{{{{{qid}}}}}"]
    changed = False
    for ph in placeholders:
        if ph in text:
            text = text.replace(ph, answer)
            changed = True
    return text, changed


def fill_docx_homework(
    input_path: str | Path,
    output_path: str | Path,
    answers: Dict[str, str],
    answer_provider: Optional[Callable[[str, str], str]] = None,
    log: Optional[Callable[[str], None]] = None,
) -> Dict[str, str]:
    doc = Document(str(input_path))
    final_answers = dict(answers)

    def logger(msg: str) -> None:
        if log:
            log(msg)

    questions = extract_questions(doc)
    logger(f"识别到题目数量: {len(questions)}")

    # 先处理全局占位符，避免题目识别不准时无法写入
    for p in _iter_paragraphs(doc):
        text = p.text
        new_text = text
        changed = False
        for qid, ans in list(final_answers.items()):
            new_text, hit = _replace_placeholders(new_text, qid, ans)
            changed = changed or hit
        if changed:
            p.text = new_text

    handled_ids = set()
    for item in questions:
        qid = item.qid
        handled_ids.add(qid)
        answer = final_answers.get(qid, "").strip()
        if not answer and answer_provider:
            logger(f"第{qid}题缺少答案，正在调用 AI 解题...")
            answer = answer_provider(qid, item.text).strip()
            final_answers[qid] = answer
            logger(f"第{qid}题 AI 已生成答案。")

        if not answer:
            logger(f"第{qid}题仍无答案，跳过。")
            continue

        # 若题目行存在占位符，直接替换
        replaced_text, changed = _replace_placeholders(item.paragraph.text, qid, answer)
        if changed:
            item.paragraph.text = replaced_text
            continue

        # 否则在题目后追加答案段落（避免覆盖原题目）
        _insert_paragraph_after(item.paragraph, f"答案：{answer}")

    # 对未识别为题目的占位符再补一遍（包括 AI 新生成答案）
    for p in _iter_paragraphs(doc):
        text = p.text
        new_text = text
        changed = False
        for qid, ans in final_answers.items():
            new_text, hit = _replace_placeholders(new_text, qid, ans)
            changed = changed or hit
        if changed:
            p.text = new_text

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return final_answers

